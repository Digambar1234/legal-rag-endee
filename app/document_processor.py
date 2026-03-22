"""
Document Processor
Handles loading, cleaning, and chunking of legal documents
(PDF, DOCX, plain-text) before they are embedded and stored in Endee.
"""

import os
import re
import uuid
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


# ------------------------------------------------------------------ #
#  Data models                                                         #
# ------------------------------------------------------------------ #

@dataclass
class DocumentChunk:
    """A single chunk of text extracted from a document."""
    chunk_id: str
    doc_id: str
    filename: str
    doc_type: str           # e.g. 'contract', 'policy', 'court_order'
    section: str            # section heading, if detected
    text: str
    char_start: int
    char_end: int
    page_number: Optional[int] = None
    extra_metadata: dict = field(default_factory=dict)

    def to_metadata(self) -> dict:
        """Serialise to flat dict for storage in Endee metadata field."""
        return {
            "chunk_id": self.chunk_id,
            "doc_id": self.doc_id,
            "filename": self.filename,
            "doc_type": self.doc_type,
            "section": self.section,
            "text": self.text,
            "char_start": self.char_start,
            "char_end": self.char_end,
            "page_number": self.page_number or 0,
            **self.extra_metadata,
        }


@dataclass
class Document:
    doc_id: str
    filename: str
    doc_type: str
    full_text: str
    pages: List[str] = field(default_factory=list)  # per-page text, if available


# ------------------------------------------------------------------ #
#  Loader                                                              #
# ------------------------------------------------------------------ #

class DocumentLoader:
    """Load documents from disk into Document objects."""

    SUPPORTED = {".pdf", ".docx", ".txt", ".md"}

    def load(self, path: str) -> Document:
        p = Path(path)
        ext = p.suffix.lower()
        if ext not in self.SUPPORTED:
            raise ValueError(f"Unsupported file type: {ext}")

        doc_id = str(uuid.uuid4())
        doc_type = self._infer_doc_type(p.stem)

        if ext == ".pdf":
            text, pages = self._load_pdf(p)
        elif ext == ".docx":
            text, pages = self._load_docx(p)
        else:
            text = p.read_text(encoding="utf-8", errors="replace")
            pages = [text]

        return Document(
            doc_id=doc_id,
            filename=p.name,
            doc_type=doc_type,
            full_text=text,
            pages=pages,
        )

    def load_directory(self, directory: str) -> List[Document]:
        docs = []
        for root, _, files in os.walk(directory):
            for f in files:
                path = os.path.join(root, f)
                if Path(path).suffix.lower() in self.SUPPORTED:
                    try:
                        docs.append(self.load(path))
                        logger.info(f"Loaded: {f}")
                    except Exception as e:
                        logger.warning(f"Skipped {f}: {e}")
        return docs

    # -- private helpers ------------------------------------------------

    def _load_pdf(self, path: Path):
        try:
            import pypdf
            pages = []
            with open(path, "rb") as fh:
                reader = pypdf.PdfReader(fh)
                for page in reader.pages:
                    pages.append(page.extract_text() or "")
            return "\n\n".join(pages), pages
        except ImportError:
            logger.warning("pypdf not installed; reading as raw bytes.")
            return path.read_text(errors="replace"), []

    def _load_docx(self, path: Path):
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            return text, [text]
        except ImportError:
            logger.warning("python-docx not installed; reading as text.")
            text = path.read_text(errors="replace")
            return text, [text]

    @staticmethod
    def _infer_doc_type(stem: str) -> str:
        stem_lower = stem.lower()
        if any(k in stem_lower for k in ["contract", "agreement", "nda", "mou"]):
            return "contract"
        if any(k in stem_lower for k in ["policy", "terms", "privacy"]):
            return "policy"
        if any(k in stem_lower for k in ["court", "judgment", "order", "ruling"]):
            return "court_order"
        if any(k in stem_lower for k in ["patent", "ip", "trademark"]):
            return "ip_document"
        return "legal_document"


# ------------------------------------------------------------------ #
#  Chunker                                                             #
# ------------------------------------------------------------------ #

class TextChunker:
    """
    Splits document text into overlapping chunks suitable for embedding.

    Strategy:
    1. Try to split on section headings (ALL-CAPS lines, numbered sections).
    2. Within each section, split on paragraph boundaries.
    3. If a paragraph exceeds max_chars, apply sliding-window character split.
    """

    HEADING_RE = re.compile(
        r"^(\d+[\.\)]\s+[A-Z]|[A-Z][A-Z\s]{5,}$|ARTICLE\s+\w+|SECTION\s+\w+)",
        re.MULTILINE,
    )

    def __init__(self, max_chars: int = 800, overlap_chars: int = 150):
        self.max_chars = max_chars
        self.overlap_chars = overlap_chars

    def chunk(self, doc: Document) -> List[DocumentChunk]:
        chunks: List[DocumentChunk] = []
        text = doc.full_text

        # Split into sections
        sections = self._split_sections(text)

        char_cursor = 0
        for section_heading, section_text in sections:
            for chunk_text, start, end in self._split_paragraphs(section_text, char_cursor):
                chunk_id = str(uuid.uuid4())
                chunks.append(
                    DocumentChunk(
                        chunk_id=chunk_id,
                        doc_id=doc.doc_id,
                        filename=doc.filename,
                        doc_type=doc.doc_type,
                        section=section_heading[:120],
                        text=chunk_text.strip(),
                        char_start=start,
                        char_end=end,
                    )
                )
            char_cursor += len(section_text)

        logger.info(f"Chunked '{doc.filename}' into {len(chunks)} chunks")
        return [c for c in chunks if len(c.text) > 30]  # drop trivially short chunks

    # -- private --------------------------------------------------------

    def _split_sections(self, text: str):
        """Return list of (heading, body) tuples."""
        positions = [m.start() for m in self.HEADING_RE.finditer(text)]
        if not positions:
            return [("General", text)]

        sections = []
        for i, pos in enumerate(positions):
            end = positions[i + 1] if i + 1 < len(positions) else len(text)
            line_end = text.find("\n", pos)
            heading = text[pos:line_end].strip() if line_end != -1 else text[pos:pos+80].strip()
            body = text[line_end + 1:end] if line_end != -1 else text[pos:end]
            sections.append((heading, body))

        # text before first heading
        if positions[0] > 0:
            sections.insert(0, ("Preamble", text[: positions[0]]))

        return sections

    def _split_paragraphs(self, text: str, base_offset: int):
        """Yield (chunk_text, abs_start, abs_end) for each chunk."""
        paragraphs = re.split(r"\n{2,}", text)
        buffer = ""
        buf_start = base_offset

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(buffer) + len(para) <= self.max_chars:
                buffer = (buffer + " " + para).strip() if buffer else para
            else:
                if buffer:
                    yield buffer, buf_start, buf_start + len(buffer)
                    # overlap: keep last overlap_chars of buffer
                    overlap = buffer[-self.overlap_chars:] if len(buffer) > self.overlap_chars else buffer
                    buf_start += len(buffer) - len(overlap)
                    buffer = overlap + " " + para
                else:
                    # single paragraph too long — slide window
                    yield from self._slide(para, base_offset)
                    buffer = ""

        if buffer:
            yield buffer, buf_start, buf_start + len(buffer)

    def _slide(self, text: str, base_offset: int):
        step = self.max_chars - self.overlap_chars
        for i in range(0, len(text), step):
            chunk = text[i: i + self.max_chars]
            yield chunk, base_offset + i, base_offset + i + len(chunk)
